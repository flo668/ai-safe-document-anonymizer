"""
Flask Anonimiseren Tool - Main Application
Multi-file anonymization tool met text/DOCX/Excel support
"""

import os
import shutil
from pathlib import Path
from datetime import datetime, timedelta
from flask import Flask
from config import config
from anonymizer import TextAnonymizer, ExcelAnonymizer, PdfAnonymizer, AnonymizationRule, ExcelColumnRule


def create_app(config_name='default'):
    """Application factory"""
    # Use custom static_url_path for VPS deployment (defaults to /static for localhost)
    static_path = os.environ.get('STATIC_URL_PATH', '/static')
    app = Flask(__name__, static_url_path=static_path)
    app.config.from_object(config[config_name])
    config[config_name].init_app(app)

    # Disable caching for development
    app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0

    # Cleanup oude bestanden bij startup
    cleanup_old_files(app)

    # Register blueprints
    from routes import register_blueprints
    register_blueprints(app)

    return app


# Helper Functions (gebruikt door blueprints)


def allowed_file(filename):
    """Check of bestand toegestaan is"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in {'txt', 'docx', 'xlsx', 'csv', 'pdf', 'md'}


def get_file_type(extension):
    """Bepaal file type van extensie"""
    ext_map = {
        'txt': 'txt',
        'md': 'txt',  # Markdown behandelen als platte tekst
        'docx': 'docx',
        'xlsx': 'xlsx',
        'csv': 'xlsx',  # CSV behandelen als Excel
        'pdf': 'pdf'
    }
    return ext_map.get(extension, 'unknown')


def get_file_preview(session_dir, file_id, phone_placeholder, email_placeholder):
    """
    Genereer preview van wat gedetecteerd zou worden in een bestand

    Args:
        session_dir: Directory met uploaded bestanden
        file_id: ID van bestand
        phone_placeholder: Placeholder voor telefoons
        email_placeholder: Placeholder voor emails

    Returns:
        dict met preview informatie
    """
    from docx import Document

    # Vind input bestand (exclude .meta.json files)
    input_file = None
    for filepath in session_dir.glob(f"{file_id}.*"):
        if not filepath.name.endswith('.meta.json'):
            input_file = filepath
            break

    if not input_file:
        raise FileNotFoundError(f"Bestand {file_id} niet gevonden")

    file_ext = input_file.suffix.lstrip('.')
    file_type = get_file_type(file_ext)

    # Alleen text/docx/pdf ondersteunen preview
    if file_type not in ['txt', 'docx', 'pdf']:
        return {
            'id': file_id,
            'fileType': file_type,
            'message': 'Preview alleen beschikbaar voor .txt, .docx en .pdf bestanden'
        }

    # Lees tekst
    if file_type == 'txt':
        with open(input_file, 'r', encoding='utf-8') as f:
            text = f.read()
    elif file_type == 'docx':
        doc = Document(input_file)
        text = '\n'.join([para.text for para in doc.paragraphs])
    elif file_type == 'pdf':
        if not PdfAnonymizer.check_support():
            return {
                'id': file_id,
                'fileType': file_type,
                'message': 'PDF ondersteuning niet beschikbaar. Installeer: pip install pdfplumber reportlab'
            }
        text = PdfAnonymizer.extract_text_for_preview(input_file, max_chars=500)
    else:
        text = ""

    # Genereer preview
    preview_data = TextAnonymizer.get_preview(
        text, phone_placeholder, email_placeholder, max_chars=500
    )

    return {
        'id': file_id,
        'fileType': file_type,
        **preview_data
    }


def process_single_file(session_dir, output_dir, file_id, rules_data, excel_rules_data, active_tab,
                        phone_placeholder='[TEL VERWIJDERD]', email_placeholder='[EMAIL VERWIJDERD]',
                        general_placeholder='[ANONIEM]', auto_detect_enabled=True,
                        reversible_mode=False, mapping=None, preserve_formatting=False,
                        sheet_names=None):
    """
    Verwerk een enkel bestand met auto-detectie en reversible mode ondersteuning

    Args:
        session_dir: Input directory
        output_dir: Output directory
        file_id: Bestand ID
        rules_data: Handmatige anonimisatie regels
        excel_rules_data: Excel kolom regels
        active_tab: Actieve tab (text/excel)
        phone_placeholder: Placeholder voor telefoonnummers
        email_placeholder: Placeholder voor emails
        general_placeholder: Standaard placeholder voor handmatige regels
        auto_detect_enabled: Of auto-detectie aan staat
        reversible_mode: Of reversible mode actief is
        mapping: AnonymizationMapping object voor reversible mode
        preserve_formatting: Of originele opmaak behouden moet worden (alleen DOCX)
        sheet_names: List van sheet names om te processen (None = alle sheets, alleen voor Excel)

    Returns:
        dict met result info en logs
    """
    from docx import Document
    import json

    # Vind input bestand (exclude .meta.json files)
    input_file = None
    for filepath in session_dir.glob(f"{file_id}.*"):
        if not filepath.name.endswith('.meta.json'):
            input_file = filepath
            break

    if not input_file:
        raise FileNotFoundError(f"Bestand {file_id} niet gevonden")

    file_ext = input_file.suffix.lstrip('.')
    file_type = get_file_type(file_ext)

    # Lees originele bestandsnaam uit metadata
    metadata_file = session_dir / f"{file_id}.meta.json"
    if metadata_file.exists():
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
            original_name = metadata.get('originalName', input_file.name)
    else:
        # Fallback naar server bestandsnaam
        original_name = input_file.name

    # Output path - formaat: {file_id}_{base_name}_ann_{HHMM}.{ext}
    timestamp = datetime.now().strftime('%H%M')  # Alleen uren en minuten

    name_parts = original_name.rsplit('.', 1)
    if len(name_parts) == 2:
        base_name, extension = name_parts
        output_filename = f"{file_id}_{base_name}_ann_{timestamp}.{extension}"
    else:
        output_filename = f"{file_id}_{original_name}_ann_{timestamp}"

    output_file = output_dir / output_filename

    logs = []
    auto_detect_report = None

    # Process based on type
    if file_type in ['txt', 'docx']:
        # Text processing MET auto-detectie
        rules = [AnonymizationRule(r) for r in rules_data]

        # Update replacement terms met general_placeholder als ze leeg zijn
        for rule in rules:
            if not rule.replacement_term and not rule.remove_instead:
                rule.replacement_term = general_placeholder

        # Process based on file type and formatting preference
        if file_type == 'txt':
            # TXT: Altijd platte tekst
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()

            anonymized_text, log_entries, auto_detect_report = TextAnonymizer.anonymize_text_with_auto_detection(
                text, rules, auto_detect_enabled, phone_placeholder, email_placeholder,
                reversible_mode, mapping
            )

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(anonymized_text)

            logs = [log.to_dict() for log in log_entries]

        else:  # docx
            if preserve_formatting:
                # Gebruik nieuwe methode die opmaak behoudt
                try:
                    log_entries, auto_detect_report = TextAnonymizer.process_docx_preserve_formatting(
                        input_file, output_file, rules, auto_detect_enabled,
                        phone_placeholder, email_placeholder, reversible_mode, mapping
                    )
                    logs = [log.to_dict() for log in log_entries]
                except Exception as e:
                    # Fallback naar oude methode als het faalt
                    print(f"Opmaak behoud gefaald, fallback naar standaard methode: {e}")
                    doc = Document(input_file)
                    text = '\n'.join([para.text for para in doc.paragraphs])

                    anonymized_text, log_entries, auto_detect_report = TextAnonymizer.anonymize_text_with_auto_detection(
                        text, rules, auto_detect_enabled, phone_placeholder, email_placeholder,
                        reversible_mode, mapping
                    )

                    new_doc = Document()
                    for line in anonymized_text.split('\n'):
                        new_doc.add_paragraph(line)
                    new_doc.save(output_file)

                    logs = [log.to_dict() for log in log_entries]
            else:
                # Oude methode: platte tekst (geen opmaak)
                doc = Document(input_file)
                text = '\n'.join([para.text for para in doc.paragraphs])

                anonymized_text, log_entries, auto_detect_report = TextAnonymizer.anonymize_text_with_auto_detection(
                    text, rules, auto_detect_enabled, phone_placeholder, email_placeholder,
                    reversible_mode, mapping
                )

                new_doc = Document()
                for line in anonymized_text.split('\n'):
                    new_doc.add_paragraph(line)
                new_doc.save(output_file)

                logs = [log.to_dict() for log in log_entries]

    elif file_type == 'pdf':
        # PDF processing MET auto-detectie
        if not PdfAnonymizer.check_support():
            raise ImportError(
                "PDF ondersteuning niet beschikbaar. "
                "Installeer: pip install pdfplumber reportlab"
            )

        rules = [AnonymizationRule(r) for r in rules_data]

        # Update replacement terms met general_placeholder als ze leeg zijn
        for rule in rules:
            if not rule.replacement_term and not rule.remove_instead:
                rule.replacement_term = general_placeholder

        # Verwerk PDF met auto-detectie
        log_entries, auto_detect_report = PdfAnonymizer.process_pdf_file(
            input_file, output_file, rules, auto_detect_enabled,
            phone_placeholder, email_placeholder
        )

        logs = [log.to_dict() for log in log_entries]

    elif file_type == 'xlsx':
        # Excel processing (met reversible mode en sheet selectie ondersteuning)
        rules = [ExcelColumnRule(r) for r in excel_rules_data]
        log_entries = ExcelAnonymizer.process_excel_file(
            input_file, output_file, rules, preserve_headers=True,
            mapping=mapping, reversible_mode=reversible_mode,
            sheet_names=sheet_names
        )
        logs = [log.to_dict() for log in log_entries]

    else:
        raise ValueError(f"Niet ondersteund bestandstype: {file_type}")

    result = {
        'id': file_id,
        'status': 'anonymized',
        'originalName': original_name,
        'anonymizedName': output_file.name,
        'logs': logs
    }

    if auto_detect_report:
        result['auto_detect_report'] = auto_detect_report

    return result


def cleanup_old_files(app):
    """Verwijder bestanden ouder dan configured tijd"""
    cutoff_time = datetime.now() - timedelta(hours=app.config['CLEANUP_OLDER_THAN_HOURS'])

    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        if not folder.exists():
            continue

        for session_dir in folder.iterdir():
            if session_dir.is_dir():
                mtime = datetime.fromtimestamp(session_dir.stat().st_mtime)
                if mtime < cutoff_time:
                    shutil.rmtree(session_dir, ignore_errors=True)


def cleanup_session_files(app, session_id):
    """Verwijder bestanden van specifieke sessie"""
    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        session_dir = folder / session_id
        if session_dir.exists():
            shutil.rmtree(session_dir, ignore_errors=True)


if __name__ == '__main__':
    # Development server (gebruik 'production' voor productie mode)
    config_name = os.environ.get('FLASK_ENV', 'development')
    app = create_app(config_name)

    # Port configuratie (standaard 5001 i.p.v. 5000 vanwege macOS AirPlay conflict)
    port = int(os.environ.get('FLASK_PORT', 5001))
    app.run(host='0.0.0.0', port=port)
