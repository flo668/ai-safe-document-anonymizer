"""
Reversible Anonymization Module
Ondersteunt anonimiseren met mapping en de-anonimiseren
"""

import json
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document


class AnonymizationMapping:
    """Beheert de mapping tussen originele en geanonimiseerde waarden"""

    def __init__(self):
        self.mappings: Dict[str, str] = {}  # {placeholder: original_value}
        self.reverse_mappings: Dict[str, str] = {}  # {original_value: placeholder}
        self.counter = 0

    def add_mapping(self, original: str, placeholder: str):
        """Voeg een mapping toe"""
        self.mappings[placeholder] = original
        self.reverse_mappings[original] = placeholder

    def get_or_create_placeholder(self, original: str, prefix: str = "[ITEM") -> str:
        """Krijg bestaande placeholder of maak nieuwe aan"""
        if original in self.reverse_mappings:
            return self.reverse_mappings[original]

        self.counter += 1
        placeholder = f"{prefix}_{self.counter}]"
        self.add_mapping(original, placeholder)
        return placeholder

    def to_dict(self) -> dict:
        """Export mapping als dictionary"""
        return {
            "mappings": self.mappings,
            "total_items": len(self.mappings),
            "version": "1.0"
        }

    def to_json(self) -> str:
        """Export mapping als JSON string"""
        return json.dumps(self.to_dict(), indent=2, ensure_ascii=False)

    @classmethod
    def from_dict(cls, data: dict) -> 'AnonymizationMapping':
        """Importeer mapping van dictionary"""
        mapping = cls()
        mapping.mappings = data.get("mappings", {})
        # Rebuild reverse mapping
        for placeholder, original in mapping.mappings.items():
            mapping.reverse_mappings[original] = placeholder
        # Set counter to max ID
        if mapping.mappings:
            ids = []
            for placeholder in mapping.mappings.keys():
                try:
                    # Extract number from "[ITEM_123]"
                    num = int(placeholder.split('_')[1].rstrip(']'))
                    ids.append(num)
                except:
                    pass
            mapping.counter = max(ids) if ids else 0
        return mapping

    @classmethod
    def from_json(cls, json_str: str) -> 'AnonymizationMapping':
        """Importeer mapping van JSON string"""
        data = json.loads(json_str)
        return cls.from_dict(data)

    @classmethod
    def from_file(cls, file_path: Path) -> 'AnonymizationMapping':
        """Laad mapping van bestand"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return cls.from_json(f.read())


class ReverseAnonymizer:
    """De-anonimisatie processor"""

    @staticmethod
    def deanonymize_text(text: str, mapping: AnonymizationMapping) -> str:
        """
        De-anonimiseer tekst met gegeven mapping

        Args:
            text: Geanonimiseerde tekst
            mapping: Mapping object

        Returns:
            De-geanonimiseerde tekst
        """
        result = text

        # Vervang alle placeholders met originele waarden
        for placeholder, original in mapping.mappings.items():
            result = result.replace(placeholder, original)

        return result

    @staticmethod
    def deanonymize_txt_file(
        input_path: Path,
        output_path: Path,
        mapping: AnonymizationMapping
    ):
        """De-anonimiseer TXT bestand"""
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()

        deanonymized = ReverseAnonymizer.deanonymize_text(text, mapping)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(deanonymized)

    @staticmethod
    def deanonymize_docx_file(
        input_path: Path,
        output_path: Path,
        mapping: AnonymizationMapping
    ):
        """De-anonimiseer DOCX bestand"""
        doc = Document(input_path)

        # Extraheer alle tekst
        full_text = '\n'.join([para.text for para in doc.paragraphs])

        # De-anonimiseer
        deanonymized = ReverseAnonymizer.deanonymize_text(full_text, mapping)

        # Maak nieuw document
        new_doc = Document()
        for line in deanonymized.split('\n'):
            new_doc.add_paragraph(line)

        new_doc.save(output_path)

    @staticmethod
    def deanonymize_excel_file(
        input_path: Path,
        output_path: Path,
        mapping: AnonymizationMapping
    ):
        """De-anonimiseer Excel bestand"""
        from openpyxl import load_workbook

        wb = load_workbook(input_path)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, (str, int, float)):
                        # Vervang placeholders (check both string and numeric values)
                        original_value = str(cell.value)
                        new_value = original_value
                        for placeholder, original in mapping.mappings.items():
                            if placeholder in new_value:
                                new_value = new_value.replace(placeholder, original)

                        # Try to convert back to number if it was originally a number
                        if new_value != original_value:
                            try:
                                cell.value = float(new_value)
                            except ValueError:
                                cell.value = new_value

        wb.save(output_path)

    @staticmethod
    def deanonymize_csv_file(
        input_path: Path,
        output_path: Path,
        mapping: AnonymizationMapping
    ):
        """De-anonimiseer CSV bestand"""
        import csv

        # Lees CSV
        with open(input_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = list(reader)

        # De-anonimiseer elke cel
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                if cell:
                    new_value = cell
                    for placeholder, original in mapping.mappings.items():
                        if placeholder in new_value:
                            new_value = new_value.replace(placeholder, original)
                    rows[i][j] = new_value

        # Schrijf terug naar CSV met UTF-8 BOM voor Excel compatibility
        # utf-8-sig adds BOM (Byte Order Mark) which signals Excel to use UTF-8
        # This prevents encoding corruption of special characters (é, ö, ü, etc.)
        with open(output_path, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.writer(f)
            writer.writerows(rows)
