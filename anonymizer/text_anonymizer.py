"""Text en DOCX anonimisatie module"""

import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.shared import Pt

try:
    from .patterns import PatternMatcher, ThreeLayerValidator, get_detection_summary
    PATTERNS_AVAILABLE = True
except ImportError:
    PATTERNS_AVAILABLE = False

try:
    from .reverse_anonymizer import AnonymizationMapping
    REVERSE_AVAILABLE = True
except ImportError:
    REVERSE_AVAILABLE = False


class AnonymizationRule:
    """Anonimisatie regel definitie"""

    def __init__(self, rule_dict: dict):
        self.id = rule_dict.get('id', '')
        self.original_term = rule_dict.get('originalTerm', '')
        self.replacement_term = rule_dict.get('replacementTerm', '')
        self.is_regex = rule_dict.get('isRegex', False)
        self.case_sensitive = rule_dict.get('caseSensitive', False)
        self.remove_instead = rule_dict.get('removeInsteadOfReplace', False)


class LogEntry:
    """Log entry voor anonimisatie acties"""

    def __init__(self, rule_id: str, original: str, pattern: str, replaced: str, count: int, source_file: str = ""):
        self.rule_id = rule_id
        self.original = original
        self.pattern = pattern
        self.replaced = replaced
        self.count = count
        self.source_file = source_file

    def to_dict(self) -> dict:
        return {
            'ruleId': self.rule_id,
            'originalTermDisplay': self.original,
            'appliedPattern': self.pattern,
            'replacedWith': self.replaced,
            'count': self.count,
            'sourceFileName': self.source_file
        }


class TextAnonymizer:
    """Text anonimisatie processor"""

    @staticmethod
    def auto_detect_patterns(text: str, confidence_threshold: float = 0.8) -> dict:
        """
        Auto-detect all PII patterns with 3-layer validation.

        Args:
            text: Text to scan
            confidence_threshold: Only return matches with confidence >= threshold

        Returns:
            {
                'matches': {
                    'phone_nl': [(match, confidence), ...],
                    'email': [...],
                    ...
                },
                'summary': {...}
            }
        """
        if not PATTERNS_AVAILABLE:
            return {
                'matches': {},
                'summary': {
                    'total_matches': 0,
                    'by_type': {},
                    'avg_confidence': 0.0,
                    'high_confidence': 0
                }
            }

        # Initialize validator
        validator = ThreeLayerValidator()

        # Run 3-layer validation
        all_results = validator.detect_all(text)

        # Filter by confidence threshold
        filtered_results = {}
        for pattern_name, matches in all_results.items():
            high_confidence = [(match, conf) for match, conf in matches if conf >= confidence_threshold]
            if high_confidence:
                filtered_results[pattern_name] = high_confidence

        # Generate summary
        summary = get_detection_summary(filtered_results)

        return {
            'matches': filtered_results,
            'summary': summary
        }

    @staticmethod
    def anonymize_text_with_auto_detection(
        text: str,
        rules: List[AnonymizationRule],
        auto_detect_enabled: bool = True,
        phone_placeholder: str = "[TEL VERWIJDERD]",
        email_placeholder: str = "[EMAIL VERWIJDERD]",
        reversible_mode: bool = False,
        mapping=None
    ) -> Tuple[str, List[LogEntry], dict]:
        """
        Anonimiseer tekst met gegeven regels EN automatische detectie

        Args:
            text: Tekst om te anonimiseren
            rules: List van AnonymizationRule objecten
            auto_detect_enabled: Of automatische detectie aan staat
            phone_placeholder: Placeholder voor telefoonnummers
            email_placeholder: Placeholder voor emails
            reversible_mode: Of reversible mode aan staat (unieke placeholders)
            mapping: AnonymizationMapping object voor reversible mode

        Returns:
            Tuple van (geanonimiseerde tekst, log entries, preview report)
        """
        current_text = text
        log_entries = []
        auto_detect_report = {"phone_numbers": {"count": 0}, "emails": {"count": 0}, "total_items": 0}

        # Eerst automatische detectie (indien enabled)
        if auto_detect_enabled and PATTERNS_AVAILABLE:
            matcher = PatternMatcher(phone_placeholder, email_placeholder)

            # Genereer preview report VOORDAT we anonimiseren
            auto_detect_report = matcher.get_preview_report(current_text)

            # Anonimiseer met patterns
            if reversible_mode and mapping is not None and REVERSE_AVAILABLE:
                # Gebruik reversible mode met unieke placeholders
                mapping_dict = {}
                current_text, phone_count, email_count, mapping_dict = matcher.anonymize_text_reversible(
                    current_text, mapping_dict
                )

                # Sla de mappings op in het AnonymizationMapping object
                for original, placeholder in mapping_dict.items():
                    mapping.add_mapping(original, placeholder)

                # Log met unieke placeholders info
                if phone_count > 0:
                    log_entries.append(LogEntry(
                        "auto_phone", "Automatische telefoon detectie (reversible)",
                        "NL Telefoon Patterns", f"{phone_count} unieke placeholder(s)", phone_count
                    ))

                if email_count > 0:
                    log_entries.append(LogEntry(
                        "auto_email", "Automatische email detectie (reversible)",
                        "Email Patterns", f"{email_count} unieke placeholder(s)", email_count
                    ))
            else:
                # Gebruik standaard mode met generieke placeholders
                current_text, phone_count, email_count = matcher.anonymize_text(current_text)

                # Log de automatische detecties
                if phone_count > 0:
                    log_entries.append(LogEntry(
                        "auto_phone", "Automatische telefoon detectie",
                        "NL Telefoon Patterns", phone_placeholder, phone_count
                    ))

                if email_count > 0:
                    log_entries.append(LogEntry(
                        "auto_email", "Automatische email detectie",
                        "Email Patterns", email_placeholder, email_count
                    ))

        # Dan handmatige regels
        manual_text, manual_logs = TextAnonymizer.anonymize_text(current_text, rules)

        return manual_text, log_entries + manual_logs, auto_detect_report

    @staticmethod
    def get_preview(
        text: str,
        phone_placeholder: str = "[TEL VERWIJDERD]",
        email_placeholder: str = "[EMAIL VERWIJDERD]",
        max_chars: int = 500
    ) -> dict:
        """
        Krijg een preview van wat automatisch gedetecteerd zou worden

        Args:
            text: Tekst om te analyseren
            phone_placeholder: Placeholder voor telefoons
            email_placeholder: Placeholder voor emails
            max_chars: Max karakters voor preview tekst

        Returns:
            Dictionary met preview informatie
        """
        if not PATTERNS_AVAILABLE:
            return {
                "preview_text": text[:max_chars],
                "phone_numbers": {"count": 0, "preview": []},
                "emails": {"count": 0, "preview": []},
                "total_items": 0
            }

        matcher = PatternMatcher(phone_placeholder, email_placeholder)
        report = matcher.get_preview_report(text)

        return {
            "preview_text": text[:max_chars],
            **report
        }

    @staticmethod
    def anonymize_text(text: str, rules: List[AnonymizationRule]) -> Tuple[str, List[LogEntry]]:
        """
        Anonimiseer tekst met gegeven regels

        Args:
            text: Tekst om te anonimiseren
            rules: List van AnonymizationRule objecten

        Returns:
            Tuple van (geanonimiseerde tekst, log entries)
        """
        current_text = text
        log_entries = []

        for rule in rules:
            if not rule.original_term and not rule.is_regex:
                continue

            if rule.is_regex and not rule.original_term:
                log_entries.append(LogEntry(
                    rule.id, rule.original_term, "(Lege Regex - Overgeslagen)",
                    "[VERWIJDERD]" if rule.remove_instead else rule.replacement_term,
                    0
                ))
                continue

            # Maak regex pattern
            pattern_string = rule.original_term
            if not rule.is_regex:
                pattern_string = re.escape(rule.original_term)

            flags = 0 if rule.case_sensitive else re.IGNORECASE

            try:
                regex = re.compile(pattern_string, flags)
            except re.error as e:
                log_entries.append(LogEntry(
                    rule.id, rule.original_term, f"FOUT: {pattern_string}",
                    "[VERWIJDERD]" if rule.remove_instead else rule.replacement_term,
                    0
                ))
                continue

            # Tel matches
            matches = regex.findall(current_text)
            match_count = len(matches)

            # Vervang in tekst
            if match_count > 0:
                replacement = "" if rule.remove_instead else rule.replacement_term
                current_text = regex.sub(replacement, current_text)

                log_entries.append(LogEntry(
                    rule.id, rule.original_term, regex.pattern,
                    "[VERWIJDERD]" if rule.remove_instead else rule.replacement_term,
                    match_count
                ))

        return current_text, log_entries

    @staticmethod
    def process_txt_file(input_path: Path, output_path: Path, rules: List[AnonymizationRule]) -> List[LogEntry]:
        """
        Verwerk een TXT bestand

        Args:
            input_path: Pad naar input bestand
            output_path: Pad naar output bestand
            rules: List van anonimisatie regels

        Returns:
            List van log entries
        """
        # Lees bestand
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()

        # Anonimiseer
        anonymized_text, log_entries = TextAnonymizer.anonymize_text(text, rules)

        # Schrijf output
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(anonymized_text)

        return log_entries

    @staticmethod
    def process_docx_file(input_path: Path, output_path: Path, rules: List[AnonymizationRule]) -> List[LogEntry]:
        """
        Verwerk een DOCX bestand (zonder opmaak behoud)

        Args:
            input_path: Pad naar input bestand
            output_path: Pad naar output bestand
            rules: List van anonimisatie regels

        Returns:
            List van log entries
        """
        # Lees DOCX bestand
        doc = Document(input_path)

        # Extraheer alle tekst
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)

        # Anonimiseer
        anonymized_text, log_entries = TextAnonymizer.anonymize_text(text, rules)

        # Maak nieuw document
        new_doc = Document()

        # Voeg geanonimiseerde tekst toe
        for line in anonymized_text.split('\n'):
            new_doc.add_paragraph(line)

        # Sla op
        new_doc.save(output_path)

        return log_entries

    @staticmethod
    def process_docx_preserve_formatting(
        input_path: Path,
        output_path: Path,
        rules: List[AnonymizationRule],
        auto_detect_enabled: bool = False,
        phone_placeholder: str = "[TEL VERWIJDERD]",
        email_placeholder: str = "[EMAIL VERWIJDERD]",
        reversible_mode: bool = False,
        mapping=None
    ) -> Tuple[List[LogEntry], dict]:
        """
        Verwerk een DOCX bestand MET opmaak behoud

        Behoudt:
        - Bold, italic, underline, font kleuren en sizes
        - Tabellen met alle opmaak
        - Afbeeldingen
        - Lijsten (bullets, nummering)

        Args:
            input_path: Pad naar input bestand
            output_path: Pad naar output bestand
            rules: List van anonimisatie regels
            auto_detect_enabled: Of auto-detectie aan staat
            phone_placeholder: Placeholder voor telefoons
            email_placeholder: Placeholder voor emails
            reversible_mode: Of reversible mode aan staat
            mapping: AnonymizationMapping object

        Returns:
            Tuple van (log entries, auto_detect report)
        """
        # Lees origineel document
        doc = Document(input_path)

        all_logs = []
        auto_detect_report = {"phone_numbers": {"count": 0}, "emails": {"count": 0}, "total_items": 0}

        # Helper functie om tekst te anonimiseren
        def anonymize_run_text(text: str) -> str:
            if not text:
                return text
            result, logs, report = TextAnonymizer.anonymize_text_with_auto_detection(
                text, rules, auto_detect_enabled, phone_placeholder,
                email_placeholder, reversible_mode, mapping
            )
            all_logs.extend(logs)
            # Update auto detect report
            if report.get('total_items', 0) > 0:
                auto_detect_report['phone_numbers']['count'] += report.get('phone_numbers', {}).get('count', 0)
                auto_detect_report['emails']['count'] += report.get('emails', {}).get('count', 0)
                auto_detect_report['total_items'] += report.get('total_items', 0)
            return result

        # Verwerk alle paragraphs (behoudt formatting per run)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = anonymize_run_text(run.text)

        # Verwerk alle tabellen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = anonymize_run_text(run.text)

        # Verwerk headers en footers in alle secties
        for section in doc.sections:
            # Header (first page, even pages, default)
            for header in [section.header, section.first_page_header, section.even_page_header]:
                for paragraph in header.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = anonymize_run_text(run.text)
                # Ook tabellen in headers checken
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if run.text:
                                        run.text = anonymize_run_text(run.text)

            # Footer (first page, even pages, default)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                for paragraph in footer.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = anonymize_run_text(run.text)
                # Ook tabellen in footers checken
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if run.text:
                                        run.text = anonymize_run_text(run.text)

        # Sla op (opmaak blijft behouden!)
        doc.save(output_path)

        return all_logs, auto_detect_report


def anonymize_text(text: str, rules_data: List[dict]) -> Tuple[str, List[dict]]:
    """
    Helper functie voor eenvoudig gebruik

    Args:
        text: Tekst om te anonimiseren
        rules_data: List van regel dictionaries

    Returns:
        Tuple van (geanonimiseerde tekst, log entries als dictionaries)
    """
    rules = [AnonymizationRule(r) for r in rules_data]
    anonymized, logs = TextAnonymizer.anonymize_text(text, rules)
    return anonymized, [log.to_dict() for log in logs]
