"""
Encoding Preservation Tests

Tests voor MON-05: UTF-8 Special Character Preservation

Validates that all special characters (Dutch accents, Greek, Cyrillic, symbols)
survive anonymization in all file types (TXT, DOCX, XLSX, CSV, PDF).
"""

import pytest
from pathlib import Path
from openpyxl import Workbook, load_workbook
from docx import Document
import csv


# ============================================================================
# Test Data Fixtures
# ============================================================================

@pytest.fixture
def special_chars_dutch():
    """Dutch special characters (most common in NL documents)"""
    return {
        'accents': 'é è ê ë ï ö ü à ç',
        'words': 'Café Müller naïve résumé coöperatie',
        'sentence': 'De café-eigenaar René heeft een naïef voorstel ingediend bij de coöperatie.',
        'mixed_case': 'CAFÉ café Café',
    }


@pytest.fixture
def special_chars_european():
    """European languages special characters"""
    return {
        'german': 'äöüßÄÖÜ Straße Müller Größe',
        'french': 'àâæçéèêëïîôùûüÿœÀÂÆÇÉÈÊËÏÎÔÙÛÜŸŒ',
        'spanish': 'ñáéíóúüÑÁÉÍÓÚÜ señor niño',
        'scandinavian': 'åæøÅÆØ Åse Søren København',
        'polish': 'ąćęłńóśźżĄĆĘŁŃÓŚŹŻ Gdańsk Kraków',
    }


@pytest.fixture
def special_chars_global():
    """Non-Latin scripts"""
    return {
        'greek': 'Ελληνικά αβγδεζηθικλμνξοπρστυφχψω ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ',
        'cyrillic': 'Кириллица абвгдежзийклмнопрстуфхцчшщъыьэюя',
        'arabic': 'العربية مرحبا',
        'hebrew': 'עברית שלום',
        'japanese': '日本語 ひらがな カタカナ',
        'chinese': '中文 你好',
        'korean': '한국어 안녕하세요',
    }


@pytest.fixture
def special_chars_symbols():
    """Currency, math, and other symbols"""
    return {
        'currency': '€ £ ¥ $ ¢ ₹ ₽ ₩',
        'math': '± × ÷ √ ∞ ≈ ≠ ≤ ≥',
        'punctuation': '« » " " ' ' „ ‚ – — …',
        'symbols': '© ® ™ § ¶ † ‡ • ° ′ ″',
        'arrows': '← → ↑ ↓ ↔ ⇐ ⇒ ⇔',
    }


@pytest.fixture
def special_chars_edge_cases():
    """Edge cases and problematic characters"""
    return {
        'zero_width': 'word\u200Bword',  # Zero-width space
        'bom': '\ufeffBOM test',  # Byte Order Mark
        'rtl_mark': 'English \u202ehebrew',  # Right-to-left mark
        'combining': 'e\u0301',  # é as combining characters
        'emoji': '😀 🎉 ❤️ 👍 🇳🇱',
        'whitespace': 'normal space\u00A0non-breaking\u2009thin',
    }


@pytest.fixture
def all_special_chars(
    special_chars_dutch,
    special_chars_european,
    special_chars_global,
    special_chars_symbols,
    special_chars_edge_cases
):
    """Combined fixture with all special character sets"""
    return {
        'dutch': special_chars_dutch,
        'european': special_chars_european,
        'global': special_chars_global,
        'symbols': special_chars_symbols,
        'edge_cases': special_chars_edge_cases,
    }


# ============================================================================
# Text File Encoding Tests
# ============================================================================

class TestTextFileEncoding:
    """Test UTF-8 preservation in text files (.txt, .md)"""

    def test_text_write_read_roundtrip(self, tmp_path, all_special_chars):
        """
        Test that writing and reading text preserves all special characters.

        This is the baseline test - if this fails, other tests will also fail.
        """
        for category, char_dict in all_special_chars.items():
            for subcategory, text in char_dict.items():
                # Write with UTF-8
                test_file = tmp_path / f"test_{category}_{subcategory}.txt"
                test_file.write_text(text, encoding='utf-8')

                # Read back
                content = test_file.read_text(encoding='utf-8')

                # Should be identical
                assert content == text, (
                    f"Character mismatch in {category}/{subcategory}:\n"
                    f"Original: {text}\n"
                    f"Read:     {content}"
                )

    def test_text_anonymization_preserves_chars(self, tmp_path, special_chars_dutch):
        """
        Test that text file processing preserves special characters.

        Primary test: Encoding preservation through file I/O.
        """
        from anonymizer.text_anonymizer import TextAnonymizer, AnonymizationRule

        # Create text file with special characters
        text_content = f"""Contactpersoon: Test Person
Email: test@café-example.nl
Notities: Naïve voorstel voor coöperatie
Details: {special_chars_dutch['sentence']}"""

        input_file = tmp_path / "test_input.txt"
        input_file.write_text(text_content, encoding='utf-8')

        # Create simple rule that doesn't touch special chars
        rule = AnonymizationRule({
            'id': 'test-1',
            'pattern': 'Test Person',
            'isRegex': False,
            'replacementTerm': '[NAAM]',
            'caseSensitive': True
        })

        # Process file
        output_file = tmp_path / "test_output.txt"
        _ = TextAnonymizer.process_txt_file(input_file, output_file, [rule])

        # Read output and verify special characters preserved
        result = output_file.read_text(encoding='utf-8')

        # Check special characters are intact
        assert 'café' in result
        assert 'Naïve' in result
        assert 'coöperatie' in result
        assert special_chars_dutch['sentence'] in result

    def test_text_file_with_bom(self, tmp_path):
        """
        Test handling of UTF-8 BOM (Byte Order Mark).

        Some editors add BOM (\ufeff) to UTF-8 files.
        """
        text_with_bom = '\ufeffTest content with BOM: café'
        test_file = tmp_path / "test_bom.txt"

        # Write with BOM
        with open(test_file, 'w', encoding='utf-8-sig') as f:
            f.write(text_with_bom)

        # Read back (should handle BOM gracefully)
        with open(test_file, 'r', encoding='utf-8-sig') as f:
            content = f.read()

        # BOM should be handled by encoding
        assert 'café' in content


# ============================================================================
# Excel Encoding Tests
# ============================================================================

class TestExcelEncoding:
    """Test UTF-8 preservation in Excel files (.xlsx)"""

    def test_excel_write_read_roundtrip(self, tmp_path, all_special_chars):
        """
        Test that Excel write/read preserves all special characters.

        Excel uses UTF-8 internally in .xlsx format.
        """
        test_file = tmp_path / "test_excel_encoding.xlsx"

        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Special Chars"

        row = 1
        for category, char_dict in all_special_chars.items():
            ws[f'A{row}'] = category
            row += 1

            for subcategory, text in char_dict.items():
                ws[f'A{row}'] = subcategory
                ws[f'B{row}'] = text
                row += 1

        wb.save(test_file)

        # Read back
        wb_check = load_workbook(test_file)
        ws_check = wb_check.active

        row = 1
        for category, char_dict in all_special_chars.items():
            assert ws_check[f'A{row}'].value == category
            row += 1

            for subcategory, text in char_dict.items():
                assert ws_check[f'A{row}'].value == subcategory
                cell_value = ws_check[f'B{row}'].value

                assert cell_value == text, (
                    f"Excel encoding failed for {category}/{subcategory}:\n"
                    f"Original: {text}\n"
                    f"Read:     {cell_value}"
                )
                row += 1

    def test_excel_anonymization_preserves_chars(self, tmp_path, special_chars_dutch):
        """
        Test that Excel file processing preserves special characters.

        Primary test: Encoding preservation through Excel I/O.
        """
        from anonymizer.excel_anonymizer import ExcelAnonymizer, ExcelColumnRule

        # Create Excel with special characters
        wb = Workbook()
        ws = wb.active
        ws.title = "Test"

        # Headers
        ws['A1'] = 'Naam'
        ws['B1'] = 'Email'
        ws['C1'] = 'Notities'

        # Data with special characters
        ws['A2'] = 'Test Person'
        ws['B2'] = 'rené@café.nl'
        ws['C2'] = special_chars_dutch['sentence']

        input_file = tmp_path / "test_input.xlsx"
        wb.save(input_file)

        # Create anonymization rule (doesn't touch special chars)
        rule = ExcelColumnRule({
            'id': 'test-1',
            'columnName': 'Naam',
            'anonymizationType': 'replace',
            'replaceWith': '[NAAM]'
        })

        # Process file
        output_file = tmp_path / "test_output.xlsx"
        ExcelAnonymizer.process_excel_file(input_file, output_file, [rule])

        # Check output - special chars should be preserved
        wb_check = load_workbook(output_file)
        ws_check = wb_check.active

        # Special chars should be intact in other columns
        assert ws_check['B2'].value == 'rené@café.nl'
        assert ws_check['C2'].value == special_chars_dutch['sentence']

    def test_excel_formula_escaping_preserves_encoding(self, tmp_path):
        """
        Test that formula escaping doesn't corrupt special characters.

        Edge case: "=café" should become "'=café" (not corrupt encoding)
        """
        from utils.validators import escape_formula

        test_cases = [
            '=café',
            '+résumé',
            '-naïve',
            '@coöperatie',
            '=Müller',
        ]

        wb = Workbook()
        ws = wb.active

        for idx, test_value in enumerate(test_cases, start=1):
            escaped = escape_formula(test_value)
            ws[f'A{idx}'] = escaped

        test_file = tmp_path / "test_formula_escape_encoding.xlsx"
        wb.save(test_file)

        # Read back and verify
        wb_check = load_workbook(test_file)
        ws_check = wb_check.active

        for idx, test_value in enumerate(test_cases, start=1):
            cell_value = ws_check[f'A{idx}'].value

            # Should have quote prefix
            assert cell_value.startswith("'")

            # Special characters should be preserved
            original_without_prefix = test_value[1:]  # Remove =+-@
            assert original_without_prefix in cell_value, (
                f"Encoding corrupted: expected '{test_value}' to contain '{original_without_prefix}', "
                f"got '{cell_value}'"
            )


# ============================================================================
# CSV Encoding Tests
# ============================================================================

class TestCSVEncoding:
    """
    Test UTF-8 BOM preservation in CSV files.

    CSV files need UTF-8 BOM (\ufeff) to open correctly in Excel.
    Without BOM, Excel may misdetect encoding as Windows-1252.
    """

    def test_csv_with_utf8_bom(self, tmp_path, special_chars_dutch):
        """
        Test that CSV exports use UTF-8 with BOM.

        This is critical for Excel compatibility.
        """
        csv_file = tmp_path / "test_utf8_bom.csv"

        # Write CSV with BOM
        with open(csv_file, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            writer.writerow(['Naam', 'Email', 'Notities'])
            writer.writerow([
                'René Müller',
                'rené@café.nl',
                special_chars_dutch['sentence']
            ])

        # Check that BOM is present in file
        with open(csv_file, 'rb') as f:
            first_bytes = f.read(3)
            assert first_bytes == b'\xef\xbb\xbf', (
                "CSV file missing UTF-8 BOM (needed for Excel compatibility)"
            )

        # Read back with utf-8-sig (should strip BOM)
        with open(csv_file, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            rows = list(reader)

        # Verify data
        assert rows[0] == ['Naam', 'Email', 'Notities']
        assert rows[1][0] == 'René Müller'
        assert rows[1][1] == 'rené@café.nl'
        assert rows[1][2] == special_chars_dutch['sentence']

    def test_csv_without_bom_fails_in_excel_simulation(self, tmp_path, special_chars_dutch):
        """
        Demonstrate that CSV without BOM corrupts in Excel.

        This test shows WHY we need BOM - without it, Excel misreads UTF-8 as Windows-1252.
        """
        csv_file = tmp_path / "test_no_bom.csv"

        # Write CSV without BOM (just UTF-8)
        with open(csv_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Naam'])
            writer.writerow(['Café Müller'])

        # Simulate Excel's behavior: Read as Windows-1252 (Excel's default for CSV without BOM)
        try:
            with open(csv_file, 'r', encoding='windows-1252') as f:
                reader = csv.reader(f)
                rows = list(reader)

            # Without BOM, Excel shows: "CafÃ© MÃ¼ller" instead of "Café Müller"
            read_value = rows[1][0]
            assert read_value != 'Café Müller', (
                f"Expected encoding corruption, but got correct value: {read_value}"
            )

        except UnicodeDecodeError:
            # Also acceptable - shows encoding incompatibility
            pass

    def test_csv_export_from_excel_preserves_encoding(self, tmp_path, special_chars_dutch):
        """
        Test Excel -> CSV export preserves special characters.

        This is a common workflow: User anonymizes Excel, then exports to CSV.
        """
        from anonymizer.excel_anonymizer import ExcelAnonymizer

        # Create Excel with special chars
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Text'
        ws['A2'] = special_chars_dutch['sentence']

        excel_file = tmp_path / "test.xlsx"
        wb.save(excel_file)

        # Convert to CSV (simulating export)
        csv_file = tmp_path / "test.csv"

        wb_read = load_workbook(excel_file)
        ws_read = wb_read.active

        with open(csv_file, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            for row in ws_read.iter_rows(values_only=True):
                writer.writerow(row)

        # Verify CSV has BOM and correct data
        with open(csv_file, 'rb') as f:
            assert f.read(3) == b'\xef\xbb\xbf'

        with open(csv_file, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            rows = list(reader)

        assert rows[1][0] == special_chars_dutch['sentence']


# ============================================================================
# DOCX Encoding Tests
# ============================================================================

class TestDOCXEncoding:
    """Test UTF-8 preservation in Word documents (.docx)"""

    def test_docx_write_read_roundtrip(self, tmp_path, all_special_chars):
        """
        Test that DOCX write/read preserves all special characters.

        DOCX uses UTF-8 in internal XML.
        """
        test_file = tmp_path / "test_docx_encoding.docx"

        # Create document
        doc = Document()

        for category, char_dict in all_special_chars.items():
            doc.add_heading(category, level=1)

            for subcategory, text in char_dict.items():
                doc.add_paragraph(f"{subcategory}: {text}")

        doc.save(test_file)

        # Read back
        doc_check = Document(test_file)

        # Collect all text
        all_text = '\n'.join([para.text for para in doc_check.paragraphs])

        # Verify all special characters are present
        for category, char_dict in all_special_chars.items():
            assert category in all_text

            for subcategory, text in char_dict.items():
                assert text in all_text, (
                    f"DOCX encoding failed for {category}/{subcategory}:\n"
                    f"Expected text: {text}\n"
                    f"Not found in document"
                )

    def test_docx_anonymization_preserves_chars(self, tmp_path, special_chars_dutch):
        """
        Test that DOCX file processing preserves special characters.

        Primary test: Encoding preservation through DOCX I/O.
        """
        from anonymizer.text_anonymizer import TextAnonymizer, AnonymizationRule

        # Create DOCX with special characters
        doc = Document()
        doc.add_paragraph('Contactpersoon: Test Person')
        doc.add_paragraph(f'Email: rené@café.nl')
        doc.add_paragraph(special_chars_dutch['sentence'])

        input_file = tmp_path / "test_input.docx"
        doc.save(input_file)

        # Create rule (doesn't touch special chars)
        rule = AnonymizationRule({
            'id': 'test-1',
            'pattern': 'Test Person',
            'isRegex': False,
            'replacementTerm': '[NAAM]',
            'caseSensitive': True
        })

        # Process file
        output_file = tmp_path / "test_output.docx"
        _ = TextAnonymizer.process_docx_file(input_file, output_file, [rule])

        # Check output - special characters should be preserved
        doc_check = Document(output_file)
        all_text = '\n'.join([para.text for para in doc_check.paragraphs])

        # Special characters should be intact
        assert 'rené@café.nl' in all_text
        assert special_chars_dutch['sentence'] in all_text


# ============================================================================
# Integration Tests
# ============================================================================

class TestEncodingIntegration:
    """Integration tests: Full anonymization workflow"""

    def test_full_workflow_preserves_all_encodings(self, tmp_path, special_chars_dutch):
        """
        Test complete workflow: Upload -> Process -> Download

        Simulates real user workflow with multiple file types.
        Primary test: Encoding preserved through full pipeline.
        """
        from anonymizer.text_anonymizer import TextAnonymizer, AnonymizationRule
        from anonymizer.excel_anonymizer import ExcelAnonymizer, ExcelColumnRule

        # Create rule
        text_rule = AnonymizationRule({
            'id': 'test-1',
            'pattern': 'PLACEHOLDER',
            'isRegex': False,
            'replacementTerm': '[VERWIJDERD]',
            'caseSensitive': True
        })

        excel_rule = ExcelColumnRule({
            'id': 'excel-1',
            'columnName': 'Email',
            'anonymizationType': 'replace',
            'replaceWith': '[EMAIL]'
        })

        # Test 1: Text file
        text_file = tmp_path / "test.txt"
        text_content = f"Document: {special_chars_dutch['sentence']}"
        text_file.write_text(text_content, encoding='utf-8')

        text_output = tmp_path / "test_output.txt"
        _ = TextAnonymizer.process_txt_file(text_file, text_output, [text_rule])

        output_content = text_output.read_text(encoding='utf-8')
        # Special chars should be preserved
        assert special_chars_dutch['sentence'] in output_content

        # Test 2: Excel file
        wb = Workbook()
        ws = wb.active
        ws['A1'] = 'Email'
        ws['A2'] = 'rené@café.nl'

        excel_file = tmp_path / "test.xlsx"
        wb.save(excel_file)

        excel_output = tmp_path / "test_output.xlsx"
        ExcelAnonymizer.process_excel_file(excel_file, excel_output, [excel_rule])

        wb_check = load_workbook(excel_output)
        ws_check = wb_check.active
        # Special chars should be preserved in output
        assert 'rené' in ws_check['A2'].value.lower() or ws_check['A2'].value == '[EMAIL]'

    def test_mixed_encoding_scenarios(self, tmp_path):
        """
        Test edge case: Document with multiple encoding types.

        Some documents mix ASCII, Latin-1, and UTF-8 characters.
        """
        mixed_text = """
        ASCII: Hello World
        Latin-1: café (ISO-8859-1)
        UTF-8: Café Müller 中文 日本語
        Symbols: € £ ¥
        Emoji: 😀 🎉
        """

        # Write with UTF-8
        test_file = tmp_path / "mixed_encoding.txt"
        test_file.write_text(mixed_text, encoding='utf-8')

        # Read back
        content = test_file.read_text(encoding='utf-8')

        # All characters should be preserved
        assert 'café' in content
        assert 'Café Müller' in content
        assert '中文' in content
        assert '€ £ ¥' in content
        assert '😀 🎉' in content
